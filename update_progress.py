import sys
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def update_docx(action_title, details_text):
    docx_path = "working_progress.docx"
    
    # Check if file exists
    if os.path.exists(docx_path):
        doc = Document(docx_path)
    else:
        doc = Document()
        # Add a title
        title = doc.add_paragraph()
        run = title.add_run("NHẬT KÝ QUÁ TRÌNH LÀM VIỆC")
        run.bold = True
        run.font.size = Pt(18)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run("Dự án: AI Video Automation Pipeline\nThư mục: d:\\New folder\n")
        sub_run.font.size = Pt(11)
        sub_run.italic = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph().add_run("-" * 50)
    
    # Add a heading for the entry
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    heading_p = doc.add_paragraph()
    heading_run = heading_p.add_run(f"[{now}] - {action_title}")
    heading_run.bold = True
    heading_run.font.size = Pt(13)
    heading_run.font.color.rgb = RGBColor(0, 102, 204) # Deep blue
    
    # Add details
    details_p = doc.add_paragraph()
    details_run = details_p.add_run(details_text)
    details_run.font.size = Pt(11)
    
    # Add a separator line
    doc.add_paragraph().add_run("-" * 40)
    
    # Save
    doc.save(docx_path)
    print(f"Updated {docx_path} successfully.")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python update_progress.py <action_title> <details_text>")
        sys.exit(1)
    
    title = sys.argv[1]
    details = sys.argv[2]
    update_docx(title, details)
